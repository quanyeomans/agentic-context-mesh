"""Unit tests for :mod:`kairix.extractors.docx` (OF-2, Wave 4).

Two seams are exercised:

  1. The **scripted-document** seam — a fake ``document_opener``
     returning an in-memory :class:`_StubDocument`. Used for shape /
     branch tests that don't need the upstream library.
  2. The **real-library** seam — invokes the actual
     :mod:`docx` (python-docx) library against the recorded fixtures
     under ``tests/fixtures/extractors/sample.docx`` and
     ``tests/fixtures/extractors/tracked_changes_sample.docx``.
     Skipped when the optional ``docx`` extra is not installed.

Sabotage-proof per test:

  * ``test_extract_preserves_heading_hierarchy`` — flipping the
    ``Heading 1`` branch of :func:`_paragraph_to_markdown` to drop
    the ``#`` prefix breaks the assertion (executed below — see
    OF-2 dispatch reporting). Restored after run.
  * ``test_extract_renders_list_items`` — removing the ``List Bullet`` /
    ``List Number`` style mapping in production drops the markers
    and breaks the assertion.
  * ``test_extract_renders_table_as_pipe_syntax`` — flipping
    :func:`_table_row_to_markdown` to omit the leading ``|`` breaks
    the assertion.
  * ``test_extract_tracked_changes_accepts_inserted_skips_deleted`` —
    flipping the track-changes walker to include ``<w:del>`` content
    (or skip ``<w:ins>``) breaks the assertion.
  * ``test_can_extract_rejects_text_mime`` — broadening
    :meth:`can_extract` (e.g. claiming ``text/*``) breaks the
    assertion.
  * ``test_quality_ok_false_when_no_headings`` — relaxing the
    heading requirement in :meth:`quality_ok` breaks the assertion.
"""

from __future__ import annotations

from collections.abc import Callable
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import pytest

from kairix.extractors import ExtractedDocument
from kairix.extractors.docx import DocxExtractor, make_extractor, version

pytestmark = pytest.mark.unit

_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

FIXTURE_DIR = Path(__file__).parent.parent / "fixtures" / "extractors"
FIXTURE_DOCX = FIXTURE_DIR / "sample.docx"
FIXTURE_TRACKED_CHANGES = FIXTURE_DIR / "tracked_changes_sample.docx"


# ---------------------------------------------------------------------------
# Scripted-document seam — in-memory stub document
# ---------------------------------------------------------------------------


@dataclass
class _StubStyle:
    name: str = "Normal"


@dataclass
class _StubParagraph:
    text: str
    style: _StubStyle = field(default_factory=_StubStyle)
    _p: Any = None


@dataclass
class _StubCell:
    text: str


@dataclass
class _StubRow:
    cells: list[_StubCell]


@dataclass
class _StubTable:
    rows: list[_StubRow]


@dataclass
class _StubBody:
    xml: str = ""


@dataclass
class _StubElement:
    body: _StubBody = field(default_factory=_StubBody)


@dataclass
class _StubCoreProperties:
    title: str | None = None
    author: str | None = None
    created: Any = None


@dataclass
class _StubDocument:
    paragraphs: list[_StubParagraph]
    tables: list[_StubTable] = field(default_factory=list)
    element: _StubElement = field(default_factory=_StubElement)
    core_properties: _StubCoreProperties = field(default_factory=_StubCoreProperties)


def _make_extractor_with_stub(doc: _StubDocument) -> DocxExtractor:
    """Build a real :class:`DocxExtractor` wired to a stub document."""

    def _opener_factory() -> Callable[[str], _StubDocument]:
        def _open(_path: str) -> _StubDocument:
            return doc

        return _open

    return DocxExtractor(version=version, document_opener=_opener_factory)


# ---------------------------------------------------------------------------
# Factory + version tests
# ---------------------------------------------------------------------------


def test_factory_returns_docx_instance() -> None:
    extractor = make_extractor()
    assert isinstance(extractor, DocxExtractor)
    assert extractor.name == "docx"
    assert extractor.version == version


def test_version_module_level_non_empty() -> None:
    """F40 sanity — module-level ``version`` is a non-empty string."""
    assert isinstance(version, str)
    assert version.strip() != ""


# ---------------------------------------------------------------------------
# can_extract dispatch tests
# ---------------------------------------------------------------------------


def test_can_extract_claims_docx_mime() -> None:
    extractor = _make_extractor_with_stub(_StubDocument(paragraphs=[]))
    assert extractor.can_extract(_DOCX_MIME, b"PK\x03\x04") is True


def test_can_extract_claims_zip_magic_when_mime_ends_with_document() -> None:
    extractor = _make_extractor_with_stub(_StubDocument(paragraphs=[]))
    # A docx served as a generic vendor mime that still ends with "document".
    assert extractor.can_extract("application/x-vendor-document", b"PK\x03\x04") is True


def test_can_extract_rejects_text_mime() -> None:
    extractor = _make_extractor_with_stub(_StubDocument(paragraphs=[]))
    assert extractor.can_extract("text/plain", b"hello") is False
    assert extractor.can_extract("text/markdown", b"# hi") is False


def test_can_extract_rejects_pdf() -> None:
    extractor = _make_extractor_with_stub(_StubDocument(paragraphs=[]))
    assert extractor.can_extract("application/pdf", b"%PDF-1.4") is False


def test_can_extract_rejects_bare_zip_without_document_hint() -> None:
    """ZIP magic alone is ambiguous; the mime hint must end with "document"."""
    extractor = _make_extractor_with_stub(_StubDocument(paragraphs=[]))
    assert extractor.can_extract("application/octet-stream", b"PK\x03\x04") is False


# ---------------------------------------------------------------------------
# Render walk — heading hierarchy
# ---------------------------------------------------------------------------


def test_extract_preserves_heading_hierarchy() -> None:
    """H1/H2/H3 paragraphs render as #/##/### markdown prefixes."""
    doc = _StubDocument(
        paragraphs=[
            _StubParagraph(text="Title", style=_StubStyle(name="Heading 1")),
            _StubParagraph(text="Section", style=_StubStyle(name="Heading 2")),
            _StubParagraph(text="Detail", style=_StubStyle(name="Heading 3")),
            _StubParagraph(text="Body paragraph."),
        ],
    )
    extractor = _make_extractor_with_stub(doc)
    extracted = extractor.extract(b"PK\x03\x04" + b"x" * 64, _DOCX_MIME)
    assert isinstance(extracted, ExtractedDocument)
    lines = extracted.markdown.splitlines()
    assert "# Title" in lines
    assert "## Section" in lines
    assert "### Detail" in lines
    assert "Body paragraph." in lines


def test_extract_renders_list_items() -> None:
    """Bullet + numbered paragraphs render as ``-`` and ``1.`` markers."""
    doc = _StubDocument(
        paragraphs=[
            _StubParagraph(text="A bullet item", style=_StubStyle(name="List Bullet")),
            _StubParagraph(text="A numbered item", style=_StubStyle(name="List Number")),
        ],
    )
    extractor = _make_extractor_with_stub(doc)
    extracted = extractor.extract(b"PK\x03\x04" + b"x" * 64, _DOCX_MIME)
    lines = extracted.markdown.splitlines()
    assert "- A bullet item" in lines
    assert "1. A numbered item" in lines


def test_extract_renders_table_as_pipe_syntax() -> None:
    """Tables render as GitHub-Flavored pipe-syntax markdown tables."""
    table = _StubTable(
        rows=[
            _StubRow(cells=[_StubCell(text="Header A"), _StubCell(text="Header B")]),
            _StubRow(cells=[_StubCell(text="Cell 1"), _StubCell(text="Cell 2")]),
        ],
    )
    doc = _StubDocument(
        paragraphs=[_StubParagraph(text="Body text long enough to register as content.")],
        tables=[table],
    )
    extractor = _make_extractor_with_stub(doc)
    extracted = extractor.extract(b"PK\x03\x04" + b"x" * 64, _DOCX_MIME)
    md = extracted.markdown
    assert "| Header A | Header B |" in md
    assert "| --- | --- |" in md
    assert "| Cell 1 | Cell 2 |" in md


def test_extract_skips_empty_paragraphs() -> None:
    """Empty paragraphs drop out of the rendered markdown."""
    doc = _StubDocument(
        paragraphs=[
            _StubParagraph(text=""),
            _StubParagraph(text="   "),
            _StubParagraph(text="Real body content here."),
        ],
    )
    extractor = _make_extractor_with_stub(doc)
    extracted = extractor.extract(b"PK\x03\x04" + b"x" * 32, _DOCX_MIME)
    # Two empty paragraphs would render as two blank double-newlines if
    # they slipped through; the markdown should start at the body text.
    assert extracted.markdown.strip() == "Real body content here."


def test_extract_metadata_pulls_title_and_author() -> None:
    """Document core_properties flow through to :class:`DocMetadata`."""
    doc = _StubDocument(
        paragraphs=[
            _StubParagraph(text="Heading", style=_StubStyle(name="Heading 1")),
            _StubParagraph(text="Body line one." * 10),
        ],
        core_properties=_StubCoreProperties(title="Sample Title", author="agent-alpha"),
    )
    extractor = _make_extractor_with_stub(doc)
    extracted = extractor.extract(b"PK\x03\x04" + b"x" * 64, _DOCX_MIME)
    assert extracted.metadata.title == "Sample Title"
    assert extracted.metadata.author == "agent-alpha"


def test_extract_returns_zero_confidence_for_empty_bytes() -> None:
    extractor = _make_extractor_with_stub(_StubDocument(paragraphs=[_StubParagraph(text="x")]))
    extracted = extractor.extract(b"", _DOCX_MIME)
    assert extracted.confidence == 0.0


def test_extract_no_tracked_changes_flag_when_body_clean() -> None:
    """The side-channel flag stays False when the body has no w:ins / w:del."""
    doc = _StubDocument(
        paragraphs=[_StubParagraph(text="Plain body without tracked changes.")],
        element=_StubElement(body=_StubBody(xml="<w:body></w:body>")),
    )
    extractor = _make_extractor_with_stub(doc)
    extractor.extract(b"PK\x03\x04" + b"x" * 32, _DOCX_MIME)
    assert extractor.last_extract_had_tracked_changes is False


# ---------------------------------------------------------------------------
# Quality gate tests
# ---------------------------------------------------------------------------


def test_quality_ok_true_when_heading_present_and_long_enough() -> None:
    doc = _StubDocument(
        paragraphs=[
            _StubParagraph(text="Section Heading", style=_StubStyle(name="Heading 1")),
            _StubParagraph(text="Body paragraph " * 16),
        ],
    )
    extractor = _make_extractor_with_stub(doc)
    extracted = extractor.extract(b"PK\x03\x04" + b"x" * 64, _DOCX_MIME)
    assert extractor.quality_ok(extracted) is True


def test_quality_ok_false_when_no_headings() -> None:
    """No heading → quality_ok is False even if the body is long enough."""
    doc = _StubDocument(
        paragraphs=[_StubParagraph(text="A long body of plain prose. " * 12)],
    )
    extractor = _make_extractor_with_stub(doc)
    extracted = extractor.extract(b"PK\x03\x04" + b"x" * 64, _DOCX_MIME)
    assert extractor.quality_ok(extracted) is False


def test_quality_ok_false_when_below_char_floor() -> None:
    doc = _StubDocument(
        paragraphs=[_StubParagraph(text="Tiny heading", style=_StubStyle(name="Heading 1"))],
    )
    extractor = _make_extractor_with_stub(doc)
    extracted = extractor.extract(b"PK\x03\x04" + b"x" * 4096, _DOCX_MIME)
    assert extractor.quality_ok(extracted) is False


# ---------------------------------------------------------------------------
# Real-library tests — exercise python-docx against the recorded fixtures
# ---------------------------------------------------------------------------


def _docx_available() -> bool:
    try:
        import docx as _docx  # noqa: F401 — probe-only import; resolved at runtime
    except ImportError:
        return False
    return True


@pytest.mark.skipif(
    not _docx_available(),
    reason="python-docx extra not installed; install via 'pip install Kairix-agentic-knowledge-mgt[docx]'",
)
def test_real_fixture_preserves_heading_hierarchy() -> None:
    """The recorded sample fixture round-trips H1/H2/H3 hierarchy intact."""
    raw = FIXTURE_DOCX.read_bytes()
    extractor = make_extractor()
    doc = extractor.extract(raw, _DOCX_MIME)
    lines = doc.markdown.splitlines()
    assert "# Introduction" in lines
    assert "## Background" in lines
    assert "### Details" in lines


@pytest.mark.skipif(
    not _docx_available(),
    reason="python-docx extra not installed; install via 'pip install Kairix-agentic-knowledge-mgt[docx]'",
)
def test_real_fixture_renders_lists_and_table() -> None:
    raw = FIXTURE_DOCX.read_bytes()
    extractor = make_extractor()
    doc = extractor.extract(raw, _DOCX_MIME)
    lines = doc.markdown.splitlines()
    assert any(line.startswith("- ") for line in lines)
    assert any(line.startswith("1. ") for line in lines)
    assert "| Header A | Header B |" in doc.markdown
    assert "| --- | --- |" in doc.markdown


@pytest.mark.skipif(
    not _docx_available(),
    reason="python-docx extra not installed; install via 'pip install Kairix-agentic-knowledge-mgt[docx]'",
)
def test_real_fixture_quality_ok_true() -> None:
    raw = FIXTURE_DOCX.read_bytes()
    extractor = make_extractor()
    doc = extractor.extract(raw, _DOCX_MIME)
    assert extractor.quality_ok(doc) is True
    # The fixture has no tracked changes; the side-channel stays False.
    assert extractor.last_extract_had_tracked_changes is False


@pytest.mark.skipif(
    not _docx_available(),
    reason="python-docx extra not installed; install via 'pip install Kairix-agentic-knowledge-mgt[docx]'",
)
def test_extract_tracked_changes_accepts_inserted_skips_deleted() -> None:
    """The tracked-changes fixture surfaces ``<w:ins>`` text and drops ``<w:del>``."""
    raw = FIXTURE_TRACKED_CHANGES.read_bytes()
    extractor = make_extractor()
    doc = extractor.extract(raw, _DOCX_MIME)
    assert "INSERTED-CONTENT" in doc.markdown
    assert "DELETED-CONTENT" not in doc.markdown
    assert extractor.last_extract_had_tracked_changes is True
