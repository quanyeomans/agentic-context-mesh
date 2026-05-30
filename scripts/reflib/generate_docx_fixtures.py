#!/usr/bin/env python3
"""Generate 10 synthetic DOCX files for the per-type fixture corpus.

ADR-028 measurement prereq. Uses python-docx (the [docx] extra). Mix of
short memos (1-2 pages) and long policy docs (15+ pages with heading
hierarchy).

Run from repo root:
    python3 scripts/reflib/generate_docx_fixtures.py

Output: reference-library/per-type-fixtures/docx/*.docx (10 files)
"""

from __future__ import annotations

import random
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt

_SCRIPT_DIR = Path(__file__).resolve().parent
sys.path.insert(0, str(_SCRIPT_DIR.parent.parent))

from scripts.reflib._fixture_vocab import (  # noqa: E402 — sys.path tweak above
    AGENTS,
    PROJECTS,
    QUARTERS,
    TOPICS,
)

SEED = 28028
N_FIXTURES = 10


def _short_memo(doc: Document, project: str, topic: str, rng: random.Random) -> None:
    """1-2 page memo: title + 2 paragraphs."""
    title = doc.add_heading(f"{project.title()} — {topic.title()}", level=1)
    for run in title.runs:
        run.font.size = Pt(16)
    doc.add_paragraph(
        f"This memo records the {topic} for {project} as of the latest cycle. "
        f"Owner: {rng.choice(AGENTS)}. Reviewers: {rng.choice(AGENTS)}, {rng.choice(AGENTS)}."
    )
    doc.add_paragraph(
        f"The {topic} unblocks the next milestone. Decision: proceed with the "
        f"plan as drafted. Next action: schedule the review for the start of "
        f"the following quarter."
    )


def _long_policy(doc: Document, project: str, rng: random.Random) -> None:
    """Long policy doc: 15+ heading-level-deep sections."""
    doc.add_heading(f"{project.title()} — Operating Policy", level=1)
    doc.add_paragraph(f"Owner: {rng.choice(AGENTS)}. Effective: this fiscal year.")
    sections = [
        "Scope and Principles",
        "Roles and Responsibilities",
        "Decision Framework",
        "Operating Cadence",
        "Reporting and Telemetry",
        "Risk Register",
        "Exceptions Process",
        "Vendor and Third-Party Engagement",
        "Security and Access Posture",
        "Incident Response",
        "Change Management",
        "Audit and Review",
        "Sunset and Retirement",
        "Glossary",
        "References",
    ]
    for sec in sections:
        doc.add_heading(sec, level=2)
        for _ in range(rng.randint(2, 4)):
            doc.add_heading(rng.choice(TOPICS).title(), level=3)
            paragraph = doc.add_paragraph()
            paragraph.add_run(
                f"This subsection covers {rng.choice(TOPICS)} for {project}. "
                f"It applies to all {rng.choice(QUARTERS)} cycles and reviews. "
                f"The owner is {rng.choice(AGENTS)}; the on-call backup is "
                f"{rng.choice(AGENTS)}. Adjust the cadence in line with the "
                f"feedback from the prior cycle. Capture changes in the audit log."
            )


def generate(output_dir: Path) -> int:
    output_dir.mkdir(parents=True, exist_ok=True)
    rng = random.Random(SEED)
    written = 0
    # 5 short memos
    for i in range(5):
        project, _focus = PROJECTS[i % len(PROJECTS)]
        topic = TOPICS[i % len(TOPICS)]
        path = output_dir / f"{project}-{topic.replace(' ', '-')}-memo.docx"
        doc = Document()
        _short_memo(doc, project, topic, rng)
        doc.save(str(path))
        written += 1
    # 5 long policies
    for i in range(5):
        project, _focus = PROJECTS[i % len(PROJECTS)]
        path = output_dir / f"{project}-operating-policy.docx"
        doc = Document()
        _long_policy(doc, project, rng)
        doc.save(str(path))
        written += 1
    return written


def main() -> int:
    repo_root = Path(__file__).resolve().parents[2]
    output_dir = repo_root / "reference-library" / "per-type-fixtures" / "docx"
    n = generate(output_dir)
    print(f"docx fixtures: wrote {n} files to {output_dir}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
